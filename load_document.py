import os
from dotenv import load_dotenv
import pinecone
from utils import LocalEmbeddings, vectorstore
from langchain_core.documents import Document
from docx import Document as DocxDocument

load_dotenv()

# Read .docx file
def read_docx(file_path):
    doc = DocxDocument(file_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():  # Only include non-empty paragraphs
            full_text.append(para.text)
    return "\n".join(full_text)

# Split text into chunks under 40KB (approx 40,000 characters, assuming 1 byte per char)
def split_text(text, max_size=40000):
    chunks = []
    current_chunk = ""
    for line in text.split("\n"):
        if len(current_chunk) + len(line) + 1 > max_size:  # +1 for newline
            chunks.append(current_chunk.strip())
            current_chunk = line
        else:
            current_chunk += "\n" + line if current_chunk else line
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks

# Path to your document
doc_path = "documents/loubby.docx"
loubby_content = read_docx(doc_path)

# Split into manageable chunks
chunks = split_text(loubby_content)
documents = [Document(page_content=chunk, metadata={"source": "loubby.docx", "chunk": i}) for i, chunk in enumerate(chunks)]

# Embed and upsert into Pinecone
embedding = LocalEmbeddings()
vectorstore.add_documents(documents, namespace="nav_indexed")

print(f"Loubby document (loubby.docx) loaded into Pinecone in {len(chunks)} chunks.")